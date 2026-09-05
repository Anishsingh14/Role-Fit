"""
Core ML pipeline for Role-Fit.
Kept separate from main.py so it can be unit-tested / imported
independently.

Pipeline stages:
    1. parse_resume_file()      -> raw text from .pdf / .docx
    2. extract_skills()         -> canonical skill list found in text
    3. build_job_vectors()      -> weighted skill matrix for all jobs (the
                                    "training data" for K-NN)
    4. vectorize_resume()       -> weighted skill vector for the resume,
                                    aligned to the same feature space
    5. ResumeJobMatcher         -> wraps sklearn NearestNeighbors (K-NN)
    6. skill_gap_report()       -> matched / missing skills + readiness score
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Tuple

import numpy as np
from sklearn.neighbors import NearestNeighbors

from job_data import MASTER_SKILLS, SKILL_SYNONYMS, get_job_dataset

# STAGE 1: FILE PARSING

def parse_resume_file(file_path: str) -> str:
    # Extract raw text from a .pdf or .docx resume file.
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Only .pdf and .docx are supported."
        )

def _parse_pdf(file_path: str) -> str:
    import pdfplumber

    text_chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    text = "\n".join(text_chunks)

    if not text.strip():
        raise ValueError(
            "Could not extract any text from the PDF. It may be a scanned "
            "image without a text layer (OCR would be required)."
        )
    return text


def _parse_docx(file_path: str) -> str:
    import docx

    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]

    # Also capture text inside tables (some resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    text = "\n".join(paragraphs)
    if not text.strip():
        raise ValueError("Could not extract any text from the DOCX file.")
    return text

# STAGE 2: SKILL EXTRACTION

def extract_skills(text: str) -> List[str]:
    """
    Scan resume text for known skills/synonyms and return the canonical
    skill names found. Case-insensitive, word-boundary aware to avoid
    partial-word false positives (e.g. 'java' inside 'javascript').
    """
    normalized = f" {text.lower()} "
    # Normalize punctuation/whitespace so multi-word phrases match reliably
    normalized = re.sub(r"[\n\r\t]+", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized)

    found = set()
    for canonical, synonyms in SKILL_SYNONYMS.items():
        for syn in synonyms:
            pattern = r"(?<![a-zA-Z0-9])" + re.escape(syn.strip()) + r"(?![a-zA-Z0-9])"
            if re.search(pattern, normalized):
                found.add(canonical)
                break
    return sorted(found)

# STAGE 3 & 4: VECTORIZATION

def build_job_vectors() -> Tuple[np.ndarray, List[str], List[dict]]:
    """
    Build the weighted skill matrix for every job in the dataset.
    Returns (job_matrix, job_titles, raw_job_dataset).
    Each row = one job, each column = one skill in MASTER_SKILLS order.
    """
    jobs = get_job_dataset()
    matrix = np.zeros((len(jobs), len(MASTER_SKILLS)))

    skill_index = {skill: i for i, skill in enumerate(MASTER_SKILLS)}

    for row, job in enumerate(jobs):
        for skill, weight in job["required_skills"].items():
            if skill in skill_index:
                matrix[row, skill_index[skill]] = weight

    titles = [job["title"] for job in jobs]
    return matrix, titles, jobs

def vectorize_resume(resume_skills: List[str]) -> np.ndarray:
    # Binary (1/0) vector for the resume, aligned to MASTER_SKILLS order.
    skill_index = {skill: i for i, skill in enumerate(MASTER_SKILLS)}
    vector = np.zeros(len(MASTER_SKILLS))
    for skill in resume_skills:
        if skill in skill_index:
            vector[skill_index[skill]] = 1.0
    return vector

# STAGE 5: K-NN MATCHER

@dataclass
class JobMatch:
    title: str
    similarity: float
    required_skills: Dict[str, int] = field(default_factory=dict)


class ResumeJobMatcher:
    """Thin wrapper around sklearn's NearestNeighbors using cosine distance,
    which behaves well on sparse, high-dimensional skill vectors."""

    def __init__(self):
        self.job_matrix, self.job_titles, self.job_dataset = build_job_vectors()
        # Binarize job matrix for cosine comparison against the resume's
        # binary vector (magnitude differences from weights are handled
        # separately in the gap analysis / readiness score).
        self.job_matrix_binary = (self.job_matrix > 0).astype(float)

        self.model = NearestNeighbors(
            n_neighbors=min(5, len(self.job_titles)),
            metric="cosine",
            algorithm="brute",
        )
        self.model.fit(self.job_matrix_binary)

    def find_top_matches(self, resume_vector: np.ndarray, k: int = 5) -> List[JobMatch]:
        k = min(k, len(self.job_titles))
        distances, indices = self.model.kneighbors(
            resume_vector.reshape(1, -1), n_neighbors=k
        )
        matches = []
        for dist, idx in zip(distances[0], indices[0]):
            similarity = 1 - dist  # cosine distance -> similarity
            job = self.job_dataset[idx]
            matches.append(
                JobMatch(
                    title=job["title"],
                    similarity=round(float(similarity), 4),
                    required_skills=job["required_skills"],
                )
            )
        return matches

# STAGE 6: SKILL GAP ANALYSIS

@dataclass
class SkillGapReport:
    job_title: str
    matched_skills: List[str]
    missing_skills: List[str]
    readiness_score: float  # 0-100, weighted by skill importance

def skill_gap_report(resume_skills: List[str], job_match: JobMatch) -> SkillGapReport:
    resume_set = set(resume_skills)
    required = job_match.required_skills

    matched = [s for s in required if s in resume_set]
    missing = [s for s in required if s not in resume_set]

    total_weight = sum(required.values()) or 1
    matched_weight = sum(w for s, w in required.items() if s in resume_set)
    readiness = round((matched_weight / total_weight) * 100, 1)

    # Sort missing skills by importance (highest weight first) so the user
    # knows what to learn first
    missing_sorted = sorted(missing, key=lambda s: -required[s])
    matched_sorted = sorted(matched, key=lambda s: -required[s])

    return SkillGapReport(
        job_title=job_match.title,
        matched_skills=matched_sorted,
        missing_skills=missing_sorted,
        readiness_score=readiness,
    )


def most_common_missing_skills(
    resume_skills: List[str], job_matches: List[JobMatch], top_n: int = 10
) -> List[Tuple[str, int]]:
    """Across all top-K matched jobs, find which missing skills appear most
    often — i.e. learning these unlocks the most opportunities."""
    resume_set = set(resume_skills)
    from collections import Counter

    counter = Counter()
    for match in job_matches:
        for skill in match.required_skills:
            if skill not in resume_set:
                counter[skill] += 1
    return counter.most_common(top_n)
